import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from N1 import extract_n1_data
from N2 import extract_n2_data
from N25 import extract_n25_data
import google.generativeai as genai
import threading
from docx import Document


genai.configure(api_key="AIzaSyAckosGZaBTj1SxNti0OyllgV80yjXWJXc")

def select_file_n1():
    file_path = filedialog.askopenfilename(title="Selecione o arquivo N1", filetypes=[("Arquivos DOCX", "*.docx")])
    if file_path:
        entry_n1.delete(0, tk.END)
        entry_n1.insert(0, file_path)

def select_file_n2():
    file_path = filedialog.askopenfilename(title="Selecione o arquivo N2", filetypes=[("Arquivos DOCX", "*.docx")])
    if file_path:
        entry_n2.delete(0, tk.END)
        entry_n2.insert(0, file_path)

def select_file_n25():
    file_path = filedialog.askopenfilename(title="Selecione o arquivo N2.5", filetypes=[("Arquivos DOCX", "*.docx")])
    if file_path:
        entry_n25.delete(0, tk.END)
        entry_n25.insert(0, file_path)

def process_files():
    n1_path = r'' + entry_n1.get()  
    n2_path = r'' + entry_n2.get()
    n25_path = r'' + entry_n25.get()
    target_orgao = entry_orgao.get().strip()  

    if not n1_path or not n2_path or not n25_path or not target_orgao:
        messagebox.showerror("Erro", "Todos os arquivos e o órgão precisam ser selecionados.")
        return

    
    n1_data = extract_n1_data(n1_path, target_orgao)
    n2_data = extract_n2_data(n2_path)
    n25_data = extract_n25_data(n25_path)

    combined_data = {
        "n1_data": n1_data,
        "n2_data": n2_data,
        "n25_data": n25_data
    }
    
    print(combined_data)


    start_loading_screen()  
    threading.Thread(target=send_to_google_ai_and_save, args=(combined_data,)).start()  # Processa em uma thread separada

def send_to_google_ai_and_save(data):
    briefing = send_to_google_ai(data)
    stop_loading_screen()
    if briefing:
        save_briefing(briefing)

def send_to_google_ai(data):
    try:
        prompt = (
            "Quero que com os dados que vão ser passados no final, você crie um modelo reunindo as informações da ata n1 e n2 e n2.5 para fazer um briefing, nesse briefing você vai fazer um contexto baseado nesse exemplo aqui: Em relação aos Compromissos de políticas públicas, foram discutidos os seguintes pontos: CG-036 - atualizações sobre o Compromisso e repactuação e desdobramento de metas e indicadores. CG-133 - Situação do Projeto ATER Agroecológico, do Plano Estadual de Produção Orgânica e desdobramentos das metas do indicador 4 (distribuição de sementes).  CG-134 - Atualização das fases e marcos.  CG-150 - Atualização do Compromisso e desdobramento das metas do indicador 3.Logo em seguida você com os dados do final, cada compromisso deve ser escrito nesse modelo aqui: CG-036 (Economia solidária, cooperativismo e associativismo) O compromisso recebeu diversas deliberações por parte do Governador na última reunião N1, sendo: O Governador definiu como meta do Compromisso a formação de cooperativas com mais de 1.000 integrantes e sugeriu utilizar a INVESTE como reforço. Portanto, a SAF deve informar como está a composição atual do total de Cooperativas, isto é, quantas atendem a meta estabelecida pelo Governador. Além disso, a Secretaria deve pontuar quais encaminhamentos têm sido dados para cumprir a meta.Ainda em relação ao ponto anterior, o Governador deliberou que o foco deve ser estruturar as cadeias do caju, ovino, caprino, peixe, leites e babaçu, que devem ter suas cooperativas fortalecidas. Desta forma, a SAF deve informar como está sendo feito o apoio para cada cadeia.A SAF agora tem o objetivo de realizar o abastecimento da CEASA com produtos da Agricultura Familiar, portanto, deve traçar um plano de ação para cumprir este objetivo.Foi determinado pelo Governador que a SAF deve verificar a viabilidade de criar uma Casa Apis na região de São Raimundo Nonato. É necessário que a SAF informe as ações para cumprir essa demanda.O Governador também definiu que a SAF deve utilizar recursos do Tesouro, ao invés de esperar recursos federais. A SAF iria apresentar os projetos à SEPLAN. Qual a situação? A SAF deve medir, principalmente, a quantidade de cooperativas e cooperados, além da produção. Desta forma, é necessário reportar se esta demanda já vem sendo cumprida."
            "ALgumas coisas tambem podem vir meio separas mas lembre-se de juntar em um unico ponto aqueles itens que não fazem sentido sozinho "
            "Aqui estão os dados dos compromissos extraídos das atas:\n\n"
            f"N1: {data['n1_data']}\n\n"
            f"N2: {data['n2_data']}\n\n"
            f"N2.5: {data['n25_data']}\n\n"
        )

        model = genai.GenerativeModel(model_name="gemini-1.5-flash")

        response = model.generate_content([prompt])

        return response.text

    except Exception as e:
        print(f"Ocorreu um erro: {e}")
        messagebox.showerror("Erro", f"Ocorreu um erro ao processar com a IA: {e}")
        return None

def save_briefing(briefing):
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Arquivo Word", "*.docx")])
    if file_path:
       
        doc = Document()

        doc.add_heading('Briefing Processado pela IA', 0)

        # Processa o texto do briefing
        lines = briefing.split("\n")
        last_meeting = None
        for line in lines:
            clean_line = line.strip().replace("**", "").replace("##", "").strip()  # Remove ## e **
            
           
            if "Reunião de Governança" in clean_line:
                current_meeting = clean_line
                if current_meeting != last_meeting:
                    doc.add_heading(clean_line, level=1)  
                    last_meeting = current_meeting  
            elif clean_line.startswith("CG-"): 
                doc.add_heading(clean_line, level=2)  
            elif clean_line:  
                doc.add_paragraph(clean_line)

        
        doc.save(file_path)
        messagebox.showinfo("Sucesso", "Briefing salvo com sucesso!")

def start_loading_screen():
    global loading_screen
    loading_screen = tk.Toplevel(root)
    loading_screen.title("Processando com IA")
    ttk.Label(loading_screen, text="Processando briefing com a IA...").pack(padx=10, pady=10)
    progress = ttk.Progressbar(loading_screen, mode="indeterminate")
    progress.pack(padx=10, pady=10)
    progress.start()

def stop_loading_screen():
    loading_screen.destroy()
    
def close_application():
    root.destroy()

root = tk.Tk()
root.title("Gerador de Briefing")

tk.Label(root, text="Arquivo N1:").grid(row=0, column=0, padx=10, pady=10)
entry_n1 = tk.Entry(root, width=50)
entry_n1.grid(row=0, column=1, padx=10, pady=10)
btn_n1 = tk.Button(root, text="Selecionar", command=select_file_n1)
btn_n1.grid(row=0, column=2, padx=10, pady=10)

tk.Label(root, text="Arquivo N2:").grid(row=1, column=0, padx=10, pady=10)
entry_n2 = tk.Entry(root, width=50)
entry_n2.grid(row=1, column=1, padx=10, pady=10)
btn_n2 = tk.Button(root, text="Selecionar", command=select_file_n2)
btn_n2.grid(row=1, column=2, padx=10, pady=10)

tk.Label(root, text="Arquivo N2.5:").grid(row=2, column=0, padx=10, pady=10)
entry_n25 = tk.Entry(root, width=50)
entry_n25.grid(row=2, column=1, padx=10, pady=10)
btn_n25 = tk.Button(root, text="Selecionar", command=select_file_n25)
btn_n25.grid(row=2, column=2, padx=10, pady=10)

tk.Label(root, text="Órgão para a N1:").grid(row=3, column=0, padx=10, pady=10)
entry_orgao = tk.Entry(root, width=50)
entry_orgao.grid(row=3, column=1, padx=10, pady=10)

btn_process = tk.Button(root, text="Gerar Briefing", command=process_files)
btn_process.grid(row=4, column=1, pady=20)

btn_close = tk.Button(root, text="Fechar Aplicação", command=close_application)
btn_close.grid(row=5, column=1, pady=10)

tk.Label(root, text="Para criar um novo briefing, feche esta janela e abra novamente.", fg="red").grid(row=6, column=1, padx=10, pady=20)


root.mainloop()

