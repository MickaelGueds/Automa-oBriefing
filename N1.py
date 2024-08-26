from docx import Document
import re

def extract_n1_data(doc_path, target_orgao):
    doc = Document(doc_path)
    commitments_data = {}
    current_commitment = None
    current_orgao = None
    current_natureza = None
    extracting_observations = False
    current_observations = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # Ignorar Statuses
        if "Status do Compromisso atual" in text:
            continue

        normalized_text = text.replace(" ", "").upper()

        # Identificar o órgão
        if "ÓRGÃO" in normalized_text:
            if ":" in text:  # Verifica se existe o caractere ':'
                current_orgao = text.split(":")[1].strip().upper()
                extracting_observations = False  # Resetar o flag de extração

        # Verificar se o órgão é o que estamos buscando
        if current_orgao == target_orgao.upper():
            # Identificar o compromisso com ou sem espaço
            if "(CG-" in normalized_text or "(CG -" in normalized_text:
                if current_commitment and (current_natureza or current_observations):
                    # Salva o compromisso anterior antes de começar um novo
                    commitments_data[current_commitment] = {
                        "natureza": current_natureza,
                        "observations": current_observations
                    }

                current_commitment = text.strip()
                commitments_data[current_commitment] = {"natureza": None, "observations": []}
                current_observations = []  # Resetar observações
                current_natureza = None  # Resetar natureza
                extracting_observations = False  # Resetar a flag de extração de observações

            # Identificar e armazenar a natureza do compromisso
            elif "NATUREZA:" in normalized_text:
                current_natureza = text.split(":", 1)[1].strip()
                # Verificar se a natureza está na linha seguinte
                if not current_natureza and i < len(doc.paragraphs) - 1:
                    next_text = doc.paragraphs[i + 1].text.strip()
                    if "Status" not in next_text:
                        current_natureza = next_text  # Captura a linha seguinte como natureza

                # Atualiza a natureza do compromisso atual
                if current_commitment:
                    commitments_data[current_commitment]["natureza"] = current_natureza

            # Iniciar extração das observações quando encontrar a palavra-chave
            elif "OBSERVAÇÕES" in normalized_text:
                extracting_observations = True
                if ":" in text:
                    observations_part = re.split(r'\t+|\s{2,}', text.split("OBSERVAÇÕES:")[1].strip())
                    current_observations.extend(observations_part)
                else:
                    current_observations.append(text)

            # Continuar coletando observações até encontrar "REUNIÃO DE GOVERNANÇA N1"
            elif extracting_observations and current_commitment:
                if "REUNIÃO DE GOVERNANÇA N1" in normalized_text:
                    extracting_observations = False  # Parar a extração ao encontrar uma nova reunião
                else:
                    current_observations.append(text.strip())

    # Captura o último compromisso se as observações estiverem ativas
    if current_commitment and (current_natureza or current_observations):
        commitments_data[current_commitment] = {
            "natureza": current_natureza,
            "observations": current_observations
        }

    return commitments_data

# Exemplo de uso
# doc_path = r'c:\Users\Mickael\Downloads\ATA N1 - GESTAO E GOVERNADORIA 04-12-2023.docx'
# target_orgao = 'SEAD'

# n1_data = extract_n1_data(doc_path, target_orgao)

# for commitment, data in n1_data.items():
#     print(f"Compromisso: {commitment}")
#     print(f"Natureza: {data['natureza']}")
#     print("Observações:")
#     if data['observations']:
#         for observation in data['observations']:
#             print(f" - {observation}")
#     else:
#         print(" - Nenhuma observação disponível.")
